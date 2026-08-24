#!/usr/bin/env python3
"""Build the camera-ready ARROW manuscript in the supplied VNICT/IEEECS layout.

The script intentionally uses the reviewed ``Làm báo (4).docx`` as its content
source.  It preserves equations and drawings, applies the 2026 template geometry,
rebuilds the two low-resolution result charts, and removes three duplicate
references while updating all in-text citation numbers in one pass.
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Polygon
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "IEEECS_CPS_2026" / "VNICT2026_Template_MSWord.docx"
OUTPUT = ROOT / "IEEECS_CPS_2026" / "ARROW_VNICT2026_IEEECS_Submission.docx"
HEADER_TEXT = (
    "Hội thảo quốc gia lần thứ XXIX: Một số vấn đề chọn lọc của "
    "Công nghệ thông tin và truyền thông – Hà Nội, 7-8/11/2026"
)


def find_source() -> Path:
    matches = sorted(ROOT.glob("*(4).docx"))
    if len(matches) != 1:
        raise RuntimeError(f"Expected one '(4).docx' source, found: {matches}")
    return matches[0]


def set_run_font(run, size_pt: float, *, bold=None, italic=None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size_pt: float, *, bold=None, italic=None) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def get_or_add_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(doc: Document) -> dict[str, object]:
    styles = {}

    normal = doc.styles["Normal"]
    set_style_font(normal, 10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 0.95
    styles["normal"] = normal

    body = get_or_add_paragraph_style(doc, "Body Text")
    set_style_font(body, 10)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Inches(0.2)
    body.paragraph_format.line_spacing = 0.95
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    styles["body"] = body

    title = get_or_add_paragraph_style(doc, "paper title")
    set_style_font(title, 16, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    styles["title"] = title

    author = get_or_add_paragraph_style(doc, "Author")
    set_style_font(author, 11)
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(0)
    author.paragraph_format.space_after = Pt(2)
    author.paragraph_format.keep_with_next = True
    styles["author"] = author

    affiliation = get_or_add_paragraph_style(doc, "Affiliation")
    set_style_font(affiliation, 10)
    affiliation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_before = Pt(0)
    affiliation.paragraph_format.space_after = Pt(0)
    affiliation.paragraph_format.keep_with_next = True
    styles["affiliation"] = affiliation

    abstract = get_or_add_paragraph_style(doc, "Abstract")
    set_style_font(abstract, 9, bold=True)
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.first_line_indent = Inches(0.2)
    abstract.paragraph_format.line_spacing = 0.95
    abstract.paragraph_format.space_before = Pt(0)
    abstract.paragraph_format.space_after = Pt(0)
    styles["abstract"] = abstract

    keywords = get_or_add_paragraph_style(doc, "key words")
    set_style_font(keywords, 9, bold=True, italic=True)
    keywords.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.first_line_indent = Inches(0.2)
    keywords.paragraph_format.line_spacing = 0.95
    keywords.paragraph_format.space_before = Pt(0)
    keywords.paragraph_format.space_after = Pt(0)
    styles["keywords"] = keywords

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 10, bold=False, italic=False)
    h1.font.small_caps = True
    h1_ppr = h1.element.get_or_add_pPr()
    h1_border = h1_ppr.find(qn("w:pBdr"))
    if h1_border is not None:
        h1_ppr.remove(h1_border)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True
    styles["h1"] = h1

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 10, bold=False, italic=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.keep_with_next = True
    styles["h2"] = h2

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 10, bold=False, italic=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.keep_with_next = True
    styles["h3"] = h3

    h5 = doc.styles["Heading 5"]
    set_style_font(h5, 10, bold=False, italic=False)
    h5.font.small_caps = True
    h5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h5.paragraph_format.first_line_indent = Pt(0)
    h5.paragraph_format.space_before = Pt(8)
    h5.paragraph_format.space_after = Pt(3)
    h5.paragraph_format.keep_with_next = True
    styles["h5"] = h5

    equation = get_or_add_paragraph_style(doc, "equation")
    set_style_font(equation, 10)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.first_line_indent = Pt(0)
    equation.paragraph_format.space_before = Pt(2)
    equation.paragraph_format.space_after = Pt(2)
    equation.paragraph_format.keep_together = True
    styles["equation"] = equation

    figure_caption = get_or_add_paragraph_style(doc, "figure caption")
    set_style_font(figure_caption, 8)
    figure_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption.paragraph_format.first_line_indent = Pt(0)
    figure_caption.paragraph_format.line_spacing = 1.0
    figure_caption.paragraph_format.space_before = Pt(1)
    figure_caption.paragraph_format.space_after = Pt(5)
    figure_caption.paragraph_format.keep_together = True
    styles["figure_caption"] = figure_caption

    table_head = get_or_add_paragraph_style(doc, "table head")
    set_style_font(table_head, 8)
    table_head.font.small_caps = True
    table_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_head.paragraph_format.first_line_indent = Pt(0)
    table_head.paragraph_format.space_before = Pt(5)
    table_head.paragraph_format.space_after = Pt(2)
    table_head.paragraph_format.keep_with_next = True
    styles["table_head"] = table_head

    references = get_or_add_paragraph_style(doc, "references")
    set_style_font(references, 8)
    references.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    references.paragraph_format.left_indent = Inches(0.16)
    references.paragraph_format.first_line_indent = Inches(-0.16)
    references.paragraph_format.line_spacing = Pt(9)
    references.paragraph_format.space_before = Pt(0)
    references.paragraph_format.space_after = Pt(0)
    styles["references"] = references

    return styles


def paragraph_num_id(paragraph) -> str | None:
    values = paragraph._p.xpath("./w:pPr/w:numPr/w:numId/@w:val")
    return values[0] if values else None


def apply_style_to_paragraph(paragraph, style, size: float, *, bold=None, italic=None):
    paragraph.style = style
    for run in paragraph.runs:
        set_run_font(run, size, bold=bold, italic=italic)


def remove_paragraph_borders(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    border = ppr.find(qn("w:pBdr"))
    if border is not None:
        ppr.remove(border)


def fix_references_and_citations(doc: Document) -> None:
    ref_heading_index = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip().upper() == "REFERENCES"
    )
    mapping = {
        **{i: i for i in range(1, 21)},
        21: 15,
        22: 21,
        23: 22,
        24: 23,
        25: 4,
        26: 24,
        27: 17,
        28: 25,
        29: 26,
        30: 27,
        31: 28,
        32: 29,
        33: 30,
    }
    citation_re = re.compile(r"\[(\d+)\]")

    for paragraph in doc.paragraphs[:ref_heading_index]:
        for text_node in paragraph._p.xpath(".//w:t"):
            if not text_node.text:
                continue

            def repl(match: re.Match[str]) -> str:
                number = int(match.group(1))
                return f"[{mapping.get(number, number)}]"

            text_node.text = citation_re.sub(repl, text_node.text)

    # The document has one trailing empty paragraph carrying the final section
    # properties; only the 33 numbered paragraphs are bibliography entries.
    references = [
        p for p in doc.paragraphs[ref_heading_index + 1 :] if paragraph_num_id(p) == "1"
    ]
    if len(references) != 33:
        raise RuntimeError(f"Expected 33 references before de-duplication, found {len(references)}")

    # Keep the complete author list of old [17] while using the fuller venue data
    # from its duplicate old [27].
    references[16].text = (
        'Z. Yuan, Y. Liu, S. Wang, Z. Cai, L. Xue, Y. Deng, J. Chen, and X. Wang, '
        '"No More Manual Tests? Evaluating and Improving ChatGPT for Unit Test '
        'Generation," Proc. ACM Softw. Eng., vol. 1, no. FSE, art. 76, 2024.'
    )

    for old_number in (27, 25, 21):
        element = references[old_number - 1]._element
        element.getparent().remove(element)


def normalize_ieee_cross_references(doc: Document) -> None:
    """Use the cross-reference forms prescribed by the IEEECS template."""
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(
            "Based on this post-hoc analysis, a strengthened assertion-oriented prompt specification"
        ):
            paragraph.text = (
                "Based on this post-hoc analysis, we revised the assertion-oriented prompt "
                "specification for subsequent ARROW runs. The revised specification requires "
                "behavior-specific assertions over returned values, object-state changes, side "
                "effects, boundary conditions, and expected exceptions. It discourages "
                "assertion-free tests, constant or tautological assertions, and generic checks "
                "such as assertNotNull when a more specific expected value can be derived from "
                "the focal-class behavior. Because this refinement was formulated after the "
                "reported experiment, the current Mutation Score results reflect the original "
                "prompts, and no empirical improvement is claimed for the revised specification."
            )
        if paragraph_num_id(paragraph) == "2":
            continue  # Keep the word "Figure" in captions.
        for text_node in paragraph._p.xpath(".//w:t"):
            if not text_node.text:
                continue
            text_node.text = re.sub(r"\bFigure ([1-4])\b", r"Fig. \1", text_node.text)
            text_node.text = text_node.text.replace("Section 3.1", "Section III-A")


def create_rq2_figure(path: Path) -> None:
    plt.rcParams.update(
        {
            "font.family": "Times New Roman",
            "font.size": 7,
            "axes.labelsize": 7,
            "xtick.labelsize": 6.5,
            "ytick.labelsize": 6.5,
            "legend.fontsize": 6.2,
        }
    )
    labels = ["No Repair", "Fixed Repair", "Adaptive Repair"]
    final_compile = [45.00, 61.67, 75.83]
    final_pass = [0.00, 38.33, 56.67]
    repair_success = [np.nan, 38.33, 56.67]
    x = np.arange(len(labels))
    width = 0.23

    fig, ax = plt.subplots(figsize=(3.0, 2.25), dpi=600)
    colors = ["#3569A8", "#E07A32", "#6C8E3D"]
    bars = [
        ax.bar(x - width, final_compile, width, label="FCSR", color=colors[0]),
        ax.bar(x, final_pass, width, label="FTPR", color=colors[1]),
        ax.bar(x + width, repair_success, width, label="RSR", color=colors[2]),
    ]
    for group in bars:
        for bar in group:
            height = bar.get_height()
            if np.isfinite(height):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    height + 1.3,
                    f"{height:.2f}",
                    ha="center",
                    va="bottom",
                    fontsize=5.4,
                    rotation=90,
                )

    ax.text(x[0] + width, 2, "n/a", ha="center", va="bottom", fontsize=5.4, rotation=90)
    ax.set_ylabel("Rate (%)")
    ax.set_ylim(0, 92)
    ax.set_xticks(x, labels)
    ax.grid(axis="y", color="#D8D8D8", linewidth=0.45)
    ax.set_axisbelow(True)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(loc="upper center", ncol=3, frameon=False, bbox_to_anchor=(0.5, 1.13))

    annotations = [
        "RA: —\nRT: —",
        "RA: 4 [3–6]\nRT: 350 [325–375] s",
        "RA: 2 [1–4]\nRT: 200 [180–220] s",
    ]
    for xi, text in zip(x, annotations):
        ax.text(
            xi,
            -0.19,
            text,
            transform=ax.get_xaxis_transform(),
            ha="center",
            va="top",
            fontsize=5.5,
            linespacing=1.15,
        )
    fig.subplots_adjust(left=0.15, right=0.99, top=0.84, bottom=0.31)
    fig.savefig(path, dpi=600, facecolor="white")
    plt.close(fig)


def create_repair_algorithm_figure(path: Path) -> None:
    """Rebuild Figure 2 as a compact, publication-resolution flow diagram."""
    plt.rcParams.update({"font.family": "Times New Roman", "font.size": 5.5})
    fig, ax = plt.subplots(figsize=(3.0, 1.55), dpi=600)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    navy = "#0B3277"
    pale = "#F6F8FC"
    edge = "#244A88"

    def box(center, number, label, width=0.15, height=0.22):
        x, y = center
        patch = FancyBboxPatch(
            (x - width / 2, y - height / 2),
            width,
            height,
            boxstyle="round,pad=0.008,rounding_size=0.012",
            facecolor=pale,
            edgecolor=edge,
            linewidth=0.8,
        )
        ax.add_patch(patch)
        circle = plt.Circle((x - width / 2 + 0.017, y + height / 2 - 0.017), 0.016, color=navy)
        ax.add_patch(circle)
        ax.text(circle.center[0], circle.center[1] - 0.001, str(number), color="white", ha="center", va="center", fontsize=4.6, fontweight="bold")
        ax.text(x, y - 0.006, label, ha="center", va="center", fontsize=5.2, linespacing=1.0)
        return patch

    def arrow(start, end, text=None, rad=0.0):
        patch = FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=5.5,
            linewidth=0.75,
            color=edge,
            connectionstyle=f"arc3,rad={rad}",
        )
        ax.add_patch(patch)
        if text:
            mx = (start[0] + end[0]) / 2
            my = (start[1] + end[1]) / 2
            ax.text(mx, my + 0.025, text, color=navy, ha="center", va="center", fontsize=4.8)

    top_y = 0.75
    top_x = [0.09, 0.29, 0.49, 0.69, 0.89]
    top_labels = ["Generate", "Validate", "Classify\nfailure", "Detect repetition\n& progress", "Select/switch\nrepair strategy"]
    for number, (x, label) in enumerate(zip(top_x, top_labels), start=1):
        box((x, top_y), number, label)
    for left, right in zip(top_x[:-1], top_x[1:]):
        arrow((left + 0.079, top_y), (right - 0.079, top_y))

    bottom_y = 0.36
    box((0.89, bottom_y), 6, "Repair")
    box((0.69, bottom_y), 7, "Re-validate")
    diamond_center = (0.47, bottom_y)
    diamond_w, diamond_h = 0.17, 0.23
    diamond = Polygon(
        [
            (diamond_center[0], diamond_center[1] + diamond_h / 2),
            (diamond_center[0] + diamond_w / 2, diamond_center[1]),
            (diamond_center[0], diamond_center[1] - diamond_h / 2),
            (diamond_center[0] - diamond_w / 2, diamond_center[1]),
        ],
        closed=True,
        facecolor=pale,
        edgecolor=edge,
        linewidth=0.8,
    )
    ax.add_patch(diamond)
    circle = plt.Circle((0.415, 0.445), 0.016, color=navy)
    ax.add_patch(circle)
    ax.text(circle.center[0], circle.center[1] - 0.001, "8", color="white", ha="center", va="center", fontsize=4.6, fontweight="bold")
    ax.text(*diamond_center, "Validation\npassed?", ha="center", va="center", fontsize=5.0, linespacing=1.0)
    box((0.25, bottom_y), 9, "Finalize\ntest suite")
    box((0.47, 0.075), 10, "Rollback", width=0.16, height=0.13)

    arrow((0.89, top_y - 0.12), (0.89, bottom_y + 0.12))
    arrow((0.81, bottom_y), (0.77, bottom_y))
    arrow((0.61, bottom_y), (0.558, bottom_y))
    arrow((0.382, bottom_y), (0.332, bottom_y), "Yes")
    arrow((0.47, bottom_y - 0.12), (0.47, 0.145), "No")
    arrow((0.55, 0.075), (0.94, 0.63), rad=-0.38)

    fig.subplots_adjust(left=0.01, right=0.99, top=0.98, bottom=0.02)
    fig.savefig(path, dpi=600, facecolor="white")
    plt.close(fig)


def create_rq3_figure(path: Path) -> None:
    plt.rcParams.update(
        {
            "font.family": "Times New Roman",
            "font.size": 6.5,
            "axes.labelsize": 7,
            "xtick.labelsize": 5.5,
            "ytick.labelsize": 5.8,
        }
    )
    configs = ["G4-ZS", "G4-FS", "G4-PA", "QW-ZS", "QW-FS", "QW-PA", "Human"]
    valid = np.array([35.0, 34.0, 68.0, 48.0, 55.0, 71.0, 100.0])
    metrics = np.array(
        [
            [85.90, 86.63, 76.70, 85.40, 54.60],
            [75.30, 78.20, 77.50, 83.70, 46.20],
            [87.50, 87.40, 77.69, 88.80, 61.00],
            [86.53, 80.90, 72.40, 84.60, 53.20],
            [80.27, 83.70, 74.80, 83.20, 56.40],
            [88.31, 88.80, 76.60, 87.30, 65.20],
            [78.36, 76.73, 80.90, 69.80, 69.10],
        ]
    ).T

    fig = plt.figure(figsize=(3.0, 3.35), dpi=600)
    grid = fig.add_gridspec(2, 1, height_ratios=[1.0, 2.0], hspace=0.42)

    ax1 = fig.add_subplot(grid[0])
    x = np.arange(len(configs))
    bars = ax1.bar(x, valid, color="#3569A8", width=0.68)
    for bar, value in zip(bars, valid):
        ax1.text(
            bar.get_x() + bar.get_width() / 2,
            value + 2,
            f"{value:.0f}",
            ha="center",
            va="bottom",
            fontsize=5.5,
        )
    ax1.set_ylabel("Valid (%)")
    ax1.set_ylim(0, 112)
    ax1.set_xticks(x, configs)
    ax1.grid(axis="y", color="#D8D8D8", linewidth=0.4)
    ax1.set_axisbelow(True)
    ax1.spines[["top", "right"]].set_visible(False)
    ax1.set_title("(a) End-to-end valid-test rate", fontsize=7, pad=3)

    ax2 = fig.add_subplot(grid[1])
    image = ax2.imshow(metrics, cmap="YlGnBu", vmin=45, vmax=90, aspect="auto")
    ax2.set_xticks(np.arange(len(configs)), configs)
    ax2.set_yticks(np.arange(5), ["IC", "LC", "BC", "MC", "MS"])
    ax2.set_title("(b) Quality conditional on valid tests (%)", fontsize=7, pad=3)
    for row in range(metrics.shape[0]):
        for col in range(metrics.shape[1]):
            value = metrics[row, col]
            color = "white" if value >= 78 else "black"
            ax2.text(col, row, f"{value:.1f}", ha="center", va="center", color=color, fontsize=5.0)
    ax2.tick_params(length=0)
    for spine in ax2.spines.values():
        spine.set_linewidth(0.45)
        spine.set_color("#666666")
    cbar = fig.colorbar(image, ax=ax2, orientation="horizontal", fraction=0.08, pad=0.14, aspect=30)
    cbar.ax.tick_params(labelsize=5.5, length=2)

    fig.subplots_adjust(left=0.16, right=0.985, top=0.96, bottom=0.08)
    fig.savefig(path, dpi=600, facecolor="white")
    plt.close(fig)


def replace_chart_images(doc: Document, algorithm_path: Path, rq2_path: Path, rq3_path: Path) -> None:
    shapes = list(doc.inline_shapes)
    if len(shapes) != 4:
        raise RuntimeError(f"Expected four inline figures, found {len(shapes)}")

    replacements = {1: algorithm_path, 2: rq2_path, 3: rq3_path}
    for index, replacement in replacements.items():
        shape = shapes[index]
        blip = shape._inline.xpath(".//a:blip")[0]
        relationship_id = blip.get(qn("r:embed"))
        image_part = doc.part.related_parts[relationship_id]
        image_part._blob = replacement.read_bytes()

    target_width = Inches(2.90)
    ratios = [1672 / 941, 3.0 / 1.55, 3.0 / 2.25, 3.0 / 3.35]
    for shape, ratio in zip(shapes, ratios):
        shape.width = target_width
        shape.height = int(target_width / ratio)


def set_cell_margins(cell, top=20, start=35, bottom=20, end=35):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    specifications = {
        "top": ("single", "8"),
        "left": ("nil", "0"),
        "bottom": ("single", "8"),
        "right": ("nil", "0"),
        "insideH": ("single", "3"),
        "insideV": ("nil", "0"),
    }
    for edge, (value, size) in specifications.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")
        borders.append(node)
    tbl_pr.append(borders)


def format_tables(doc: Document) -> None:
    widths_by_table = [
        [760, 780, 1900, 768],
        [1250, 660, 766, 766, 766],
        [900, 980, 550, 760, 1018],
    ]
    for table_index, (table, widths) in enumerate(zip(doc.tables, widths_by_table)):
        if len(widths) != len(table.columns):
            raise RuntimeError(f"Unexpected column count in table {table_index + 1}")
        total = sum(widths)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table_pr = table._tbl.tblPr
        table_width = table_pr.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_pr.insert(0, table_width)
        table_width.set(qn("w:w"), str(total))
        table_width.set(qn("w:type"), "dxa")

        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        for grid_column, width in zip(grid_columns, widths):
            grid_column.set(qn("w:w"), str(width))

        set_table_borders(table)
        header_row_pr = table.rows[0]._tr.get_or_add_trPr()
        header_repeat = OxmlElement("w:tblHeader")
        header_repeat.set(qn("w:val"), "true")
        header_row_pr.append(header_repeat)

        for row_index, row in enumerate(table.rows):
            row_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            row_pr.append(cant_split)
            for col_index, (cell, width) in enumerate(zip(row.cells, widths)):
                cell.width = Twips(width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_width.set(qn("w:w"), str(width))
                tc_width.set(qn("w:type"), "dxa")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 0.9
                    paragraph.paragraph_format.keep_together = True
                    if row_index == 0 or col_index >= 2:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, 8, bold=(row_index == 0))


def set_columns(section, count: int, gap_twips: int = 360) -> None:
    sect_pr = section._sectPr
    columns = sect_pr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        sect_pr.append(columns)
    for child in list(columns):
        columns.remove(child)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), str(gap_twips))
    columns.set(qn("w:equalWidth"), "1")


def configure_section(section, *, first: bool) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Twips(11907)
    section.page_height = Twips(16840)
    section.left_margin = Twips(1559)
    section.right_margin = Twips(1559)
    section.header_distance = Twips(964)
    section.footer_distance = Twips(2268)
    if first:
        section.top_margin = Twips(1531)
        section.bottom_margin = Twips(3005)
        set_columns(section, 1)
    else:
        section.top_margin = Twips(1701)
        section.bottom_margin = Twips(2835)
        section.start_type = WD_SECTION_START.CONTINUOUS
        set_columns(section, 2, 360)


def clear_extra_paragraphs(container) -> None:
    paragraphs = container.paragraphs
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def configure_header_footer(doc: Document) -> None:
    first = doc.sections[0]
    first.header.is_linked_to_previous = False
    header = first.header
    clear_extra_paragraphs(header)
    paragraph = header.paragraphs[0]
    paragraph.text = HEADER_TEXT
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 8.5, italic=True)

    footer = first.footer
    clear_extra_paragraphs(footer)
    footer.paragraphs[0].text = ""

    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True


def format_document(doc: Document, styles: dict[str, object]) -> None:
    paragraphs = doc.paragraphs
    if len(paragraphs) < 180:
        raise RuntimeError("The source manuscript structure is not the expected reviewed version")

    fixed = {
        0: (styles["title"], 16, True, False),
        1: (styles["author"], 11, False, False),
        2: (styles["affiliation"], 10, False, False),
        3: (styles["affiliation"], 10, False, False),
        4: (styles["affiliation"], 9, False, False),
        5: (styles["abstract"], 9, True, False),
        6: (styles["keywords"], 9, True, True),
    }
    for index, (style, size, bold, italic) in fixed.items():
        apply_style_to_paragraph(paragraphs[index], style, size, bold=bold, italic=italic)

    reference_mode = False
    for index, paragraph in enumerate(paragraphs[7:], start=7):
        if paragraph.text.strip().upper() == "REFERENCES":
            reference_mode = True
            apply_style_to_paragraph(paragraph, styles["h5"], 10, bold=False, italic=False)
            continue
        if reference_mode:
            if paragraph_num_id(paragraph) == "1":
                apply_style_to_paragraph(paragraph, styles["references"], 8)
            else:
                paragraph.style = styles["normal"]
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            continue

        style_id = paragraph.style.style_id if paragraph.style is not None else ""
        if style_id == "Heading1":
            apply_style_to_paragraph(paragraph, styles["h1"], 10, bold=False, italic=False)
            remove_paragraph_borders(paragraph)
        elif style_id == "Heading2":
            apply_style_to_paragraph(paragraph, styles["h2"], 10, bold=False, italic=True)
        elif style_id == "Heading3":
            apply_style_to_paragraph(paragraph, styles["h3"], 10, bold=False, italic=True)
        elif paragraph_num_id(paragraph) == "2":
            apply_style_to_paragraph(paragraph, styles["figure_caption"], 8)
        elif paragraph_num_id(paragraph) == "3":
            apply_style_to_paragraph(paragraph, styles["table_head"], 8)
        elif paragraph._p.xpath(".//m:oMath | .//m:oMathPara") and not paragraph.text.strip():
            apply_style_to_paragraph(paragraph, styles["equation"], 10)
        elif paragraph._p.xpath(".//w:drawing"):
            paragraph.style = styles["normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        else:
            apply_style_to_paragraph(paragraph, styles["body"], 10)

    # Notes below Table II and Figure 4 use the compact note sizes prescribed by
    # the template.  They are intentionally not treated as ordinary body text.
    for note_index, size in ((123, 6.5), (164, 6.5)):
        paragraph = paragraphs[note_index]
        paragraph.style = styles["normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.9
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            set_run_font(run, size)

    # Keep each figure caption with its figure; Figure 4 also has an explanatory
    # abbreviation note immediately after the caption.
    for image_index in (34, 65, 150, 162):
        paragraphs[image_index].paragraph_format.keep_with_next = True
    paragraphs[163].paragraph_format.keep_with_next = True


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    track = settings.find(qn("w:trackRevisions"))
    if track is not None:
        settings.remove(track)


def validate_output(path: Path) -> None:
    check = Document(path)
    if len(check.sections) != 2:
        raise RuntimeError(f"Expected two sections, found {len(check.sections)}")
    if len(check.inline_shapes) != 4 or len(check.tables) != 3:
        raise RuntimeError("Figure/table count changed while building the submission")

    ref_heading = next(i for i, p in enumerate(check.paragraphs) if p.text.strip() == "REFERENCES")
    refs = [p for p in check.paragraphs[ref_heading + 1 :] if paragraph_num_id(p) == "1"]
    if len(refs) != 30:
        raise RuntimeError(f"Expected 30 unique references, found {len(refs)}")

    citations = set()
    for paragraph in check.paragraphs[:ref_heading]:
        citations.update(int(n) for n in re.findall(r"\[(\d+)\]", paragraph.text))
    if citations != set(range(1, 31)):
        raise RuntimeError(f"Citation coverage is not 1..30: {sorted(citations)}")

    if check.sections[1]._sectPr.xpath("./w:cols/@w:num") != ["2"]:
        raise RuntimeError("Main manuscript section is not two-column")
    if check.sections[1]._sectPr.xpath("./w:cols/@w:space") != ["360"]:
        raise RuntimeError("Main manuscript column gap does not match the template")

    with tempfile.TemporaryDirectory(prefix="arrow-output-media-") as temp_name:
        temp = Path(temp_name)
        from zipfile import ZipFile

        with ZipFile(path) as archive:
            image_names = [n for n in archive.namelist() if n.startswith("word/media/")]
            for name in image_names:
                destination = temp / Path(name).name
                destination.write_bytes(archive.read(name))
        dimensions = []
        for image_path in sorted(temp.iterdir()):
            try:
                with Image.open(image_path) as image:
                    dimensions.append((image_path.name, image.size))
            except Exception:
                continue
        if max(width for _, (width, _) in dimensions) < 1000:
            raise RuntimeError("No high-resolution media were retained in the output")

    print(f"Created: {path}")
    print(f"Paragraphs: {len(check.paragraphs)}; references: {len(refs)}")
    print(f"Tables: {len(check.tables)}; figures: {len(check.inline_shapes)}")
    print(f"Media dimensions: {dimensions}")


def main() -> None:
    source = find_source()
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)

    doc = Document(source)
    fix_references_and_citations(doc)
    normalize_ieee_cross_references(doc)
    styles = configure_styles(doc)
    format_document(doc, styles)

    for index, section in enumerate(doc.sections):
        configure_section(section, first=(index == 0))
    configure_header_footer(doc)
    format_tables(doc)
    enable_field_updates(doc)

    with tempfile.TemporaryDirectory(prefix="arrow-ieeecs-build-") as temp_name:
        temp = Path(temp_name)
        rq2_path = temp / "rq2_repair_results_600dpi.png"
        rq3_path = temp / "rq3_quality_results_600dpi.png"
        algorithm_path = temp / "adaptive_repair_algorithm_600dpi.png"
        create_repair_algorithm_figure(algorithm_path)
        create_rq2_figure(rq2_path)
        create_rq3_figure(rq3_path)
        replace_chart_images(doc, algorithm_path, rq2_path, rq3_path)

        doc.core_properties.title = (
            "ARROW: An Adaptive Repository-Aware Repair Workflow for LLM-Based Java Unit Test Generation"
        )
        doc.core_properties.subject = "VNICT 2026 camera-ready manuscript"
        doc.core_properties.author = (
            "Nguyen Tuan Dung; Vuong Kieu Anh; Doan Huu Minh Quang; Ngo Duc Chinh; "
            "Tran Van Han; Nguyen Hoang An; Nguyen Thi Nguyet"
        )
        OUTPUT.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUTPUT)

    validate_output(OUTPUT)


if __name__ == "__main__":
    main()
